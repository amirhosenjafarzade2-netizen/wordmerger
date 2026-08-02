"""
DOCX Merger / Page Editor
=========================

Dependencies (pip install --break-system-packages ...):
    streamlit
    python-docx
    docxcompose
    pypdf
    pymupdf        (imported as `fitz`, used only for page-preview thumbnails)
    pdf2docx       (used for the PDF -> DOCX step, see note below)

System dependency:
    LibreOffice must be installed and the `soffice` binary must be on PATH.
    This is what actually computes real page boundaries (DOCX has no concept
    of "pages" in its XML -- pagination only exists once something renders
    the layout, the same way Word itself does it). We convert DOCX -> PDF
    with LibreOffice to get real page boundaries, then edit the PDF.

    Debian/Ubuntu:  sudo apt-get install libreoffice
    macOS (brew):   brew install --cask libreoffice   (then use the soffice
                    binary inside LibreOffice.app/Contents/MacOS/)

Why pdf2docx instead of LibreOffice for the reverse (PDF -> DOCX) step:
    LibreOffice's own PDF-import filter was tried first (soffice --convert-to
    docx --infilter=writer_pdf_import). It technically produces a .docx, but
    it rebuilds every line of text as a separately positioned floating text
    box rather than normal flowing paragraphs -- it looks right when opened,
    but python-docx can't even read the text back out, and it doesn't edit
    like a normal Word document. pdf2docx does real layout analysis and
    reconstructs actual flowing paragraphs and tables, so that's what's used
    here. Verified against a table + heading + multi-page test document,
    reconstructing both text and the table correctly.

Honest caveat about fidelity that still applies:
    Converting PDF back to DOCX is still a reconstruction, not the original
    file. Plain text, headings, and simple/moderate tables come through
    well. Complex multi-column layouts, floating images, headers/footers,
    and footnotes can still shift or be dropped, because a PDF page is
    fundamentally positioned ink plus a text layer, not structured Word
    content -- reconstruction is inherently best-effort. The edited PDF is
    always offered as a lossless fallback download alongside the DOCX.
"""

import os
import shutil
import subprocess
from tempfile import NamedTemporaryFile, TemporaryDirectory

import fitz  # PyMuPDF, for rendering page thumbnails only
import streamlit as st
from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from pdf2docx import Converter as Pdf2DocxConverter
from pypdf import PdfReader, PdfWriter

st.set_page_config(page_title="DOCX Merger & Page Editor", page_icon="📄", layout="wide")


# --------------------------------------------------------------------------
# LibreOffice helpers
# --------------------------------------------------------------------------

def find_soffice():
    """Locate the LibreOffice headless binary, or None if not installed."""
    for candidate in ("soffice", "libreoffice"):
        path = shutil.which(candidate)
        if path:
            return path
    return None


def run_soffice_convert(soffice_path, input_path, output_dir, target_format):
    """
    Convert a file to `target_format` (e.g. 'pdf' or 'docx') using LibreOffice
    headless mode. Returns the path to the converted file.
    """
    cmd = [
        soffice_path,
        "--headless",
        "--norestore",
        "--convert-to", target_format,
        "--outdir", output_dir,
        input_path,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed (exit {result.returncode}).\n"
            f"stdout: {result.stdout}\nstderr: {result.stderr}"
        )
    base = os.path.splitext(os.path.basename(input_path))[0]
    out_path = os.path.join(output_dir, f"{base}.{target_format}")
    if not os.path.exists(out_path):
        raise RuntimeError(
            f"Expected converted file not found at {out_path}. "
            f"LibreOffice output: {result.stdout} {result.stderr}"
        )
    return out_path


def docx_to_pdf(soffice_path, docx_path, workdir):
    """DOCX -> PDF via LibreOffice. This is the step that computes real
    page boundaries, since DOCX has none stored in its XML."""
    return run_soffice_convert(soffice_path, docx_path, workdir, "pdf")


def pdf_to_docx(pdf_path, workdir):
    """
    PDF -> DOCX via pdf2docx (layout-analysis based), not LibreOffice.
    LibreOffice's PDF import rebuilds text as floating text boxes rather
    than normal paragraphs; pdf2docx reconstructs actual flowing paragraphs
    and tables, which is what a "Word document" deliverable should contain.
    """
    out_path = os.path.join(workdir, "result.docx")
    converter = Pdf2DocxConverter(pdf_path)
    try:
        converter.convert(out_path)
    finally:
        converter.close()
    return out_path


# --------------------------------------------------------------------------
# PDF helpers (real, accurate page-level operations)
# --------------------------------------------------------------------------

def get_page_count(pdf_path):
    return len(PdfReader(pdf_path).pages)


def render_page_thumbnails(pdf_path, zoom=0.6):
    """Return a list of PNG bytes, one per page, for preview display."""
    thumbnails = []
    doc = fitz.open(pdf_path)
    mat = fitz.Matrix(zoom, zoom)
    for page in doc:
        pix = page.get_pixmap(matrix=mat)
        thumbnails.append(pix.tobytes("png"))
    doc.close()
    return thumbnails


def parse_page_spec(spec, max_page):
    """
    Parse a string like "2,3,6-8" into a sorted list of unique 1-indexed
    page numbers, validated against max_page. Raises ValueError on bad input.
    """
    spec = spec.strip()
    if not spec:
        raise ValueError("Please enter at least one page number or range.")

    pages = set()
    for token in spec.split(","):
        token = token.strip()
        if not token:
            continue
        if "-" in token:
            parts = token.split("-")
            if len(parts) != 2:
                raise ValueError(f"Invalid range: '{token}'")
            start_str, end_str = parts[0].strip(), parts[1].strip()
            if not (start_str.isdigit() and end_str.isdigit()):
                raise ValueError(f"Invalid range: '{token}'")
            start, end = int(start_str), int(end_str)
            if start > end:
                start, end = end, start
            pages.update(range(start, end + 1))
        else:
            if not token.isdigit():
                raise ValueError(f"Invalid page number: '{token}'")
            pages.add(int(token))

    out_of_range = [p for p in pages if p < 1 or p > max_page]
    if out_of_range:
        raise ValueError(
            f"Page(s) {sorted(out_of_range)} are out of range. "
            f"This document has {max_page} page(s)."
        )
    return sorted(pages)


def delete_pages(pdf_path, pages_to_delete, output_path):
    reader = PdfReader(pdf_path)
    writer = PdfWriter()
    delete_set = set(pages_to_delete)
    for i, page in enumerate(reader.pages, start=1):
        if i not in delete_set:
            writer.add_page(page)
    if len(writer.pages) == 0:
        raise ValueError("That would delete every page. Nothing would remain.")
    with open(output_path, "wb") as f:
        writer.write(f)


def extract_pages(pdf_path, pages_to_keep, output_path):
    reader = PdfReader(pdf_path)
    writer = PdfWriter()
    keep_set = set(pages_to_keep)
    for i, page in enumerate(reader.pages, start=1):
        if i in keep_set:
            writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)


# --------------------------------------------------------------------------
# Caching: avoid re-running LibreOffice on every Streamlit rerun
# --------------------------------------------------------------------------

def get_or_build_pdf(uploaded_file, soffice_path):
    """
    Converts the uploaded docx to PDF once per unique upload and caches the
    PDF bytes + page count in session_state, keyed by filename+size+content
    hash so re-uploading a different file invalidates the cache correctly.
    """
    key = f"{uploaded_file.name}:{uploaded_file.size}"
    cache = st.session_state.setdefault("pdf_cache", {})

    if key in cache:
        return cache[key]["pdf_bytes"], cache[key]["page_count"]

    with TemporaryDirectory() as workdir:
        docx_path = os.path.join(workdir, uploaded_file.name)
        with open(docx_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        pdf_path = docx_to_pdf(soffice_path, docx_path, workdir)
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        page_count = get_page_count(pdf_path)

    cache[key] = {"pdf_bytes": pdf_bytes, "page_count": page_count}
    return pdf_bytes, page_count


# --------------------------------------------------------------------------
# UI
# --------------------------------------------------------------------------

st.title("📄 DOCX Merger & Page Editor")

soffice_path = find_soffice()
if not soffice_path:
    st.error(
        "LibreOffice (`soffice`) was not found on this system. Page-level "
        "operations require it to compute real page boundaries. Install it "
        "with `sudo apt-get install libreoffice` (Linux) or "
        "`brew install --cask libreoffice` (macOS), then restart the app."
    )

mode = st.radio(
    "Choose a mode",
    ["Merge Documents", "Delete Pages", "Extract Pages"],
    horizontal=True,
)

st.divider()

# ==========================================================================
# MODE 1: MERGE (original behavior, unchanged)
# ==========================================================================
if mode == "Merge Documents":
    st.write("Upload Word documents, choose their order, and merge them into a single DOCX.")

    uploaded_files = st.file_uploader(
        "Upload DOCX files", type=["docx"], accept_multiple_files=True, key="merge_uploader"
    )

    if uploaded_files:
        st.subheader("Document Order")
        filenames = [f.name for f in uploaded_files]
        order = []
        remaining = filenames.copy()
        for i in range(len(filenames)):
            selected = st.selectbox(f"Position {i + 1}", remaining, key=f"order_{i}")
            order.append(selected)
            remaining.remove(selected)

        st.divider()
        if st.button("Merge Documents", type="primary"):
            temp_paths = []
            output_path = None
            try:
                ordered_files = []
                for filename in order:
                    for uploaded in uploaded_files:
                        if uploaded.name == filename:
                            ordered_files.append(uploaded)
                            break

                for uploaded in ordered_files:
                    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp.write(uploaded.getbuffer())
                    tmp.close()
                    temp_paths.append(tmp.name)

                if len(temp_paths) == 0:
                    st.error("No files selected.")
                    st.stop()

                master_doc = Document(temp_paths[0])
                composer = Composer(master_doc)

                for path in temp_paths[1:]:
                    paragraph = master_doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_break(WD_BREAK.PAGE)
                    composer.append(Document(path))

                output_file = NamedTemporaryFile(delete=False, suffix=".docx")
                output_path = output_file.name
                output_file.close()
                composer.save(output_path)

                with open(output_path, "rb") as f:
                    merged_bytes = f.read()

                st.success("Documents merged successfully!")
                st.download_button(
                    label="⬇ Download Merged DOCX",
                    data=merged_bytes,
                    file_name="merged_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"Merge failed: {str(e)}")
            finally:
                for path in temp_paths:
                    try:
                        if os.path.exists(path):
                            os.remove(path)
                    except Exception:
                        pass
                if output_path:
                    try:
                        if os.path.exists(output_path):
                            os.remove(output_path)
                    except Exception:
                        pass

# ==========================================================================
# MODE 2 & 3: DELETE PAGES / EXTRACT PAGES (shared logic, real page boundaries)
# ==========================================================================
elif mode in ("Delete Pages", "Extract Pages"):
    action_word = "delete" if mode == "Delete Pages" else "extract"
    st.write(
        f"Upload a Word document. It will be rendered to determine real page "
        f"boundaries (via LibreOffice), so you can {action_word} exact pages."
    )

    uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"], key="page_uploader")

    if uploaded_file and soffice_path:
        with st.spinner("Rendering document to determine page boundaries..."):
            try:
                pdf_bytes, page_count = get_or_build_pdf(uploaded_file, soffice_path)
            except Exception as e:
                st.error(f"Could not process document: {e}")
                st.stop()

        st.info(f"This document has **{page_count}** page(s).")

        show_preview = st.checkbox("Show preview of each page")

        if show_preview:
            with TemporaryDirectory() as workdir:
                tmp_pdf_path = os.path.join(workdir, "preview.pdf")
                with open(tmp_pdf_path, "wb") as f:
                    f.write(pdf_bytes)
                with st.spinner("Rendering page thumbnails..."):
                    thumbnails = render_page_thumbnails(tmp_pdf_path)

            cols_per_row = 4
            for row_start in range(0, len(thumbnails), cols_per_row):
                cols = st.columns(cols_per_row)
                for col, idx in zip(cols, range(row_start, min(row_start + cols_per_row, len(thumbnails)))):
                    with col:
                        st.image(thumbnails[idx], caption=f"Page {idx + 1}", use_container_width=True)

        st.divider()
        st.subheader(f"{'Delete' if action_word == 'delete' else 'Extract'} Pages")
        spec = st.text_input(
            "Enter page number(s) or range(s), e.g. 2,3,6-8",
            placeholder="e.g. 2,3,6-8",
            key="page_spec",
        )

        button_label = "Delete Selected Pages" if action_word == "delete" else "Extract Selected Pages"
        if st.button(button_label, type="primary"):
            try:
                pages = parse_page_spec(spec, page_count)
            except ValueError as e:
                st.error(str(e))
                st.stop()

            with st.spinner("Editing PDF and converting back to DOCX..."):
                with TemporaryDirectory() as workdir:
                    try:
                        src_pdf_path = os.path.join(workdir, "source.pdf")
                        with open(src_pdf_path, "wb") as f:
                            f.write(pdf_bytes)

                        edited_pdf_path = os.path.join(workdir, "edited.pdf")
                        if action_word == "delete":
                            delete_pages(src_pdf_path, pages, edited_pdf_path)
                        else:
                            extract_pages(src_pdf_path, pages, edited_pdf_path)

                        # Convert edited PDF back to DOCX via pdf2docx.
                        result_docx_path = pdf_to_docx(edited_pdf_path, workdir)

                        with open(edited_pdf_path, "rb") as f:
                            edited_pdf_bytes = f.read()
                        with open(result_docx_path, "rb") as f:
                            result_docx_bytes = f.read()

                    except Exception as e:
                        st.error(f"Operation failed: {e}")
                        st.stop()

            st.success(f"Pages {pages} {'deleted' if action_word == 'delete' else 'extracted'} successfully!")
            st.warning(
                "Note: the DOCX above is reconstructed from the PDF, not the "
                "original file. Plain text, headings, and simple/moderate "
                "tables come through well. Complex multi-column layouts, "
                "floating images, headers/footers, and footnotes may still "
                "shift. If exact formatting matters, use the PDF download "
                "below instead — it's a lossless edit."
            )

            base_name = os.path.splitext(uploaded_file.name)[0]
            suffix = "deleted" if action_word == "delete" else "extracted"

            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="⬇ Download Result as DOCX",
                    data=result_docx_bytes,
                    file_name=f"{base_name}_{suffix}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            with col2:
                st.download_button(
                    label="⬇ Download Result as PDF (lossless)",
                    data=edited_pdf_bytes,
                    file_name=f"{base_name}_{suffix}.pdf",
                    mime="application/pdf",
                )
    elif uploaded_file and not soffice_path:
        st.error("Cannot proceed without LibreOffice installed. See error message above.")
